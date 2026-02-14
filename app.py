"""
SRT to DOCX Converter - Web Application
Deploy: streamlit run app.py
"""

import streamlit as st
import re
import os
import io
import zipfile
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


# ═══════════════════════════════════════════════════
# PAGE CONFIGURATION
# ═══════════════════════════════════════════════════
st.set_page_config(
    page_title="SRT to DOCX Converter",
    page_icon="📝",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ═══════════════════════════════════════════════════
# CUSTOM CSS STYLING
# ═══════════════════════════════════════════════════
st.markdown("""
<style>
    /* Main container */
    .main .block-container {
        padding-top: 2rem;
        padding-bottom: 2rem;
        max-width: 1200px;
    }

    /* Header styling */
    .main-header {
        text-align: center;
        padding: 1.5rem 0;
        background: linear-gradient(135deg, #1a478a 0%, #5a9bd5 100%);
        border-radius: 15px;
        margin-bottom: 2rem;
        color: white;
    }
    .main-header h1 {
        color: white !important;
        font-size: 2.5rem !important;
        margin-bottom: 0.3rem !important;
    }
    .main-header p {
        color: #e0e0e0;
        font-size: 1.1rem;
        margin: 0;
    }

    /* Card styling */
    .card {
        background: white;
        padding: 1.5rem;
        border-radius: 12px;
        box-shadow: 0 2px 12px rgba(0,0,0,0.08);
        margin-bottom: 1rem;
        border-left: 4px solid #1a478a;
    }

    /* File info styling */
    .file-info {
        background: #f8f9fa;
        padding: 0.8rem 1.2rem;
        border-radius: 8px;
        margin: 0.3rem 0;
        border-left: 3px solid #5a9bd5;
        font-family: 'Consolas', monospace;
        font-size: 0.9rem;
    }

    /* Success message */
    .success-box {
        background: #d4edda;
        border: 1px solid #c3e6cb;
        border-radius: 10px;
        padding: 1.2rem;
        text-align: center;
        margin: 1rem 0;
    }

    /* Stats boxes */
    .stat-box {
        background: linear-gradient(135deg, #f8f9fa 0%, #e9ecef 100%);
        padding: 1rem;
        border-radius: 10px;
        text-align: center;
        border: 1px solid #dee2e6;
    }
    .stat-box h3 {
        margin: 0;
        color: #1a478a;
        font-size: 1.8rem;
    }
    .stat-box p {
        margin: 0;
        color: #666;
        font-size: 0.85rem;
    }

    /* Hide streamlit branding */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}

    /* Upload area */
    .stFileUploader > div > div {
        border-radius: 12px;
    }

    /* Button styling */
    .stDownloadButton > button {
        width: 100%;
        border-radius: 10px;
        padding: 0.7rem 2rem;
        font-weight: bold;
        font-size: 1.1rem;
    }
</style>
""", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════
# SRT PARSER
# ═══════════════════════════════════════════════════
class SRTParser:
    """Parse SRT subtitle files."""

    BLOCK_PATTERN = re.compile(
        r'(\d+)\s*\n'
        r'(\d{2}:\d{2}:\d{2}[,\.]\d{3})\s*-->\s*(\d{2}:\d{2}:\d{2}[,\.]\d{3})\s*\n'
        r'((?:(?!\d+\s*\n\d{2}:\d{2}:\d{2}).+\n?)+)',
        re.MULTILINE
    )
    TIMESTAMP_PATTERN = re.compile(
        r'(\d{2}:\d{2}:\d{2}[,\.]\d{3})\s*-->\s*(\d{2}:\d{2}:\d{2}[,\.]\d{3})'
    )
    TAG_PATTERN = re.compile(r'<[^>]+>|{[^}]+}')

    ENCODINGS = ['utf-8-sig', 'utf-8', 'latin-1', 'cp1252', 'iso-8859-1']

    @classmethod
    def parse(cls, content_bytes):
        """Parse SRT content from bytes."""
        content = cls._decode(content_bytes)
        content = content.replace('\r\n', '\n').replace('\r', '\n').strip()
        content = content.lstrip('\ufeff')

        subtitles = cls._regex_parse(content)
        if not subtitles:
            subtitles = cls._block_parse(content)

        return subtitles

    @classmethod
    def _decode(cls, content_bytes):
        """Try multiple encodings to decode bytes."""
        for enc in cls.ENCODINGS:
            try:
                return content_bytes.decode(enc)
            except (UnicodeDecodeError, UnicodeError):
                continue
        return content_bytes.decode('latin-1', errors='replace')

    @classmethod
    def _clean_text(cls, text):
        """Remove HTML/ASS tags from subtitle text."""
        text = cls.TAG_PATTERN.sub('', text)
        lines = [l.strip() for l in text.split('\n') if l.strip()]
        return '\n'.join(lines)

    @classmethod
    def _regex_parse(cls, content):
        """Parse using regex."""
        subtitles = []
        for match in cls.BLOCK_PATTERN.findall(content):
            try:
                subtitles.append({
                    'index': int(match[0]),
                    'start_time': match[1].replace(',', '.'),
                    'end_time': match[2].replace(',', '.'),
                    'text': cls._clean_text(match[3])
                })
            except (ValueError, IndexError):
                continue
        return subtitles

    @classmethod
    def _block_parse(cls, content):
        """Fallback block-based parsing."""
        subtitles = []
        for block in content.split('\n\n'):
            block = block.strip()
            if not block:
                continue

            lines = block.split('\n')
            if len(lines) < 3:
                continue

            try:
                index = int(lines[0].strip())
            except ValueError:
                continue

            time_match = cls.TIMESTAMP_PATTERN.match(lines[1].strip())
            if not time_match:
                continue

            text = cls._clean_text('\n'.join(lines[2:]))
            if text:
                subtitles.append({
                    'index': index,
                    'start_time': time_match.group(1).replace(',', '.'),
                    'end_time': time_match.group(2).replace(',', '.'),
                    'text': text
                })

        return subtitles


# ═══════════════════════════════════════════════════
# DOCX CREATOR
# ═══════════════════════════════════════════════════
class DOCXCreator:
    """Create DOCX documents from subtitle data."""

    # Colors
    PRIMARY = RGBColor(0x1A, 0x47, 0x8A)
    SECONDARY = RGBColor(0x5A, 0x9B, 0xD5)
    TEXT = RGBColor(0x33, 0x33, 0x33)
    MUTED = RGBColor(0x88, 0x88, 0x88)
    LIGHT = RGBColor(0xCC, 0xCC, 0xCC)
    TIMESTAMP = RGBColor(0x7F, 0x8C, 0x8D)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)

    @classmethod
    def create(cls, subtitles, filename, style="table"):
        """Create DOCX and return as bytes."""
        doc = Document()

        # Page setup
        section = doc.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

        # Title
        title = doc.add_heading(level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(f"📝 Subtitles: {filename}")
        run.font.size = Pt(18)
        run.font.color.rgb = cls.PRIMARY

        # Metadata
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = meta.add_run(
            f"Converted: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | "
            f"Subtitles: {len(subtitles)} | Style: {style.title()}"
        )
        r.font.size = Pt(8)
        r.font.color.rgb = cls.MUTED
        r.italic = True

        doc.add_paragraph("")

        # Content based on style
        style_map = {
            'table': cls._table,
            'plain': cls._plain,
            'formatted': cls._formatted,
            'text_only': cls._text_only,
            'script': cls._script,
        }
        style_map[style](doc, subtitles)

        # Footer
        doc.add_paragraph("")
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = footer.add_run("Generated by SRT to DOCX Converter")
        fr.font.size = Pt(7)
        fr.font.color.rgb = cls.MUTED

        # Save to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    @classmethod
    def _table(cls, doc, subtitles):
        """Table format."""
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        table.autofit = True

        headers = ['#', 'Start', 'End', 'Subtitle Text']
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(h)
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = cls.WHITE
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for sub in subtitles:
            cells = table.add_row().cells

            # Index
            p = cells[0].paragraphs[0]
            r = p.add_run(str(sub['index']))
            r.font.size = Pt(9)
            r.font.color.rgb = cls.PRIMARY
            r.bold = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Start time
            p = cells[1].paragraphs[0]
            r = p.add_run(sub['start_time'])
            r.font.size = Pt(8)
            r.font.name = 'Consolas'
            r.font.color.rgb = cls.TIMESTAMP
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # End time
            p = cells[2].paragraphs[0]
            r = p.add_run(sub['end_time'])
            r.font.size = Pt(8)
            r.font.name = 'Consolas'
            r.font.color.rgb = cls.TIMESTAMP
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Text
            p = cells[3].paragraphs[0]
            r = p.add_run(sub['text'])
            r.font.size = Pt(10)
            r.font.color.rgb = cls.TEXT

        for row in table.rows:
            row.cells[0].width = Cm(1.2)
            row.cells[1].width = Cm(3.2)
            row.cells[2].width = Cm(3.2)
            row.cells[3].width = Cm(10.5)

    @classmethod
    def _plain(cls, doc, subtitles):
        """Plain format with separators."""
        for i, sub in enumerate(subtitles):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)

            r = p.add_run(f"[{sub['index']}] ")
            r.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = cls.PRIMARY

            r = p.add_run(f"⏱ {sub['start_time']} → {sub['end_time']}")
            r.font.size = Pt(8)
            r.font.color.rgb = cls.TIMESTAMP
            r.italic = True

            tp = doc.add_paragraph()
            tp.paragraph_format.left_indent = Inches(0.3)
            tp.paragraph_format.space_after = Pt(4)
            r = tp.add_run(sub['text'])
            r.font.size = Pt(11)
            r.font.color.rgb = cls.TEXT

            if i < len(subtitles) - 1:
                sp = doc.add_paragraph()
                r = sp.add_run("─" * 65)
                r.font.size = Pt(5)
                r.font.color.rgb = cls.LIGHT

    @classmethod
    def _formatted(cls, doc, subtitles):
        """Formatted inline style."""
        for sub in subtitles:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)

            r = p.add_run(f"[{sub['start_time']} – {sub['end_time']}]  ")
            r.font.size = Pt(8)
            r.font.color.rgb = cls.TIMESTAMP
            r.italic = True
            r.font.name = 'Consolas'

            r = p.add_run(sub['text'])
            r.font.size = Pt(11)
            r.font.color.rgb = cls.TEXT

    @classmethod
    def _text_only(cls, doc, subtitles):
        """Text only - no timestamps."""
        for sub in subtitles:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            r = p.add_run(sub['text'])
            r.font.size = Pt(12)
            r.font.name = 'Georgia'
            r.font.color.rgb = cls.TEXT

    @classmethod
    def _script(cls, doc, subtitles):
        """Screenplay style."""
        for sub in subtitles:
            tp = doc.add_paragraph()
            tp.paragraph_format.space_before = Pt(10)
            tp.paragraph_format.space_after = Pt(2)
            r = tp.add_run(f"[{sub['start_time']} → {sub['end_time']}]")
            r.font.size = Pt(8)
            r.font.color.rgb = cls.SECONDARY
            r.bold = True
            r.font.name = 'Courier New'

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.left_indent = Inches(1.5)
            p.paragraph_format.right_indent = Inches(1.5)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(sub['text'])
            r.font.size = Pt(11)
            r.font.name = 'Courier New'
            r.font.color.rgb = cls.TEXT


# ═══════════════════════════════════════════════════
# HELPER FUNCTIONS
# ═══════════════════════════════════════════════════
def format_size(size_bytes):
    """Format bytes to human readable."""
    if size_bytes < 1024:
        return f"{size_bytes} B"
    elif size_bytes < 1024 * 1024:
        return f"{size_bytes / 1024:.1f} KB"
    return f"{size_bytes / (1024 * 1024):.1f} MB"


def create_zip(files_dict):
    """Create a ZIP file from dict of {filename: bytes_buffer}."""
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        for name, buffer in files_dict.items():
            zf.writestr(name, buffer.getvalue())
    zip_buffer.seek(0)
    return zip_buffer


# ═══════════════════════════════════════════════════
# MAIN APPLICATION UI
# ═══════════════════════════════════════════════════

# ── Header ──
st.markdown("""
<div class="main-header">
    <h1>📝 SRT to DOCX Converter</h1>
    <p>Upload subtitle files and download formatted Word documents instantly</p>
</div>
""", unsafe_allow_html=True)

# ── Sidebar Settings ──
with st.sidebar:
    st.markdown("## ⚙️ Settings")
    st.markdown("---")

    # Format selection
    st.markdown("### 📄 Output Format")
    format_style = st.radio(
        "Choose document style:",
        options=['table', 'plain', 'formatted', 'text_only', 'script'],
        format_func=lambda x: {
            'table': '📊 Table — Structured columns',
            'plain': '📄 Plain — Numbered with separators',
            'formatted': '🎨 Formatted — Inline timestamps',
            'text_only': '📝 Text Only — No timestamps',
            'script': '🎬 Script — Screenplay style',
        }[x],
        index=0,
        key="format_style"
    )

    st.markdown("---")

    # Format preview descriptions
    st.markdown("### 📖 Format Preview")
    descriptions = {
        'table': "Clean table with **#**, **Start Time**, **End Time**, and **Text** columns.",
        'plain': "Each subtitle numbered with **timestamps** and horizontal **separators**.",
        'formatted': "Paragraphs with **small inline timestamps** before each subtitle text.",
        'text_only': "**Only the subtitle text** — no numbers, no timestamps. Clean reading.",
        'script': "**Screenplay style** — timestamps as scene markers, centered dialogue text.",
    }
    st.info(descriptions[format_style])

    st.markdown("---")

    # Help section
    with st.expander("❓ How to use"):
        st.markdown("""
        1. **Upload** one or more `.srt` files
        2. **Choose** your preferred format style
        3. **Preview** the parsed subtitles
        4. **Download** individual DOCX files or all as ZIP
        """)

    with st.expander("📋 What is an SRT file?"):
        st.markdown("""
        SRT (SubRip Text) is the most common subtitle format.

        ```
        1
        00:00:01,000 --> 00:00:04,500
        Hello, welcome!

        2
        00:00:05,000 --> 00:00:08,200
        This is a subtitle file.
        ```
        """)

# ── Main Content Area ──
col_upload, col_spacer, col_results = st.columns([5, 0.5, 5])

# ── LEFT COLUMN: Upload ──
with col_upload:
    st.markdown("### 📁 Upload SRT Files")

    uploaded_files = st.file_uploader(
        "Drag and drop SRT files here",
        type=['srt'],
        accept_multiple_files=True,
        help="Upload one or more .srt subtitle files"
    )

    if uploaded_files:
        st.markdown(f"### 📋 Files Uploaded ({len(uploaded_files)})")

        total_size = 0
        for uf in uploaded_files:
            size = len(uf.getvalue())
            total_size += size
            st.markdown(
                f'<div class="file-info">📄 {uf.name} &nbsp;&nbsp;'
                f'<span style="color:#888">({format_size(size)})</span></div>',
                unsafe_allow_html=True
            )

        # Stats row
        st.markdown("")
        c1, c2, c3 = st.columns(3)
        with c1:
            st.markdown(
                f'<div class="stat-box"><h3>{len(uploaded_files)}</h3>'
                f'<p>Files Uploaded</p></div>',
                unsafe_allow_html=True
            )
        with c2:
            st.markdown(
                f'<div class="stat-box"><h3>{format_size(total_size)}</h3>'
                f'<p>Total Size</p></div>',
                unsafe_allow_html=True
            )
        with c3:
            st.markdown(
                f'<div class="stat-box"><h3>{format_style.title()}</h3>'
                f'<p>Output Format</p></div>',
                unsafe_allow_html=True
            )

# ── RIGHT COLUMN: Results ──
with col_results:
    if uploaded_files:
        st.markdown("### 📥 Converted Files")

        # Process all files
        converted_files = {}
        all_stats = []

        progress_bar = st.progress(0)
        status_text = st.empty()

        for idx, uploaded_file in enumerate(uploaded_files):
            filename = uploaded_file.name
            base_name = os.path.splitext(filename)[0]
            docx_name = f"{base_name}.docx"

            status_text.text(f"⏳ Converting: {filename}...")

            try:
                # Parse SRT
                content_bytes = uploaded_file.getvalue()
                subtitles = SRTParser.parse(content_bytes)

                if not subtitles:
                    st.warning(f"⚠️ **{filename}**: No subtitles found. Skipped.")
                    all_stats.append({
                        'file': filename, 'status': '⚠️ Skipped',
                        'count': 0
                    })
                    continue

                # Create DOCX
                docx_buffer = DOCXCreator.create(
                    subtitles, filename, style=format_style
                )
                converted_files[docx_name] = docx_buffer

                all_stats.append({
                    'file': filename,
                    'status': '✅ Success',
                    'count': len(subtitles)
                })

            except Exception as e:
                st.error(f"❌ **{filename}**: Error — {str(e)}")
                all_stats.append({
                    'file': filename, 'status': '❌ Failed',
                    'count': 0
                })

            # Update progress
            progress_bar.progress((idx + 1) / len(uploaded_files))

        status_text.text("✅ Conversion complete!")

        # ── Download Section ──
        if converted_files:
            st.markdown("")
            st.markdown(
                '<div class="success-box">'
                f'<h3>🎉 {len(converted_files)} file(s) converted successfully!</h3>'
                '</div>',
                unsafe_allow_html=True
            )

            st.markdown("---")

            # Individual download buttons
            st.markdown("#### 📄 Download Individual Files")

            for docx_name, docx_buffer in converted_files.items():
                c1, c2 = st.columns([3, 1])
                with c1:
                    st.markdown(f"📄 **{docx_name}**")
                with c2:
                    st.download_button(
                        label="⬇️ Download",
                        data=docx_buffer,
                        file_name=docx_name,
                        mime="application/vnd.openxmlformats-officedocument"
                             ".wordprocessingml.document",
                        key=f"dl_{docx_name}"
                    )

            # Download ALL as ZIP
            if len(converted_files) > 1:
                st.markdown("---")
                st.markdown("#### 📦 Download All Files")

                zip_buffer = create_zip(converted_files)

                st.download_button(
                    label=f"📦 Download All ({len(converted_files)} files) as ZIP",
                    data=zip_buffer,
                    file_name="converted_subtitles.zip",
                    mime="application/zip",
                    key="dl_zip"
                )

            # ── Conversion Report ──
            st.markdown("---")
            st.markdown("#### 📊 Conversion Report")

            for stat in all_stats:
                st.markdown(
                    f"- {stat['status']} **{stat['file']}** "
                    f"— {stat['count']} subtitles"
                )

        # ── Preview Section ──
        st.markdown("---")
        st.markdown("### 👁️ Preview Subtitles")

        if uploaded_files:
            preview_file = st.selectbox(
                "Select file to preview:",
                options=[f.name for f in uploaded_files],
                key="preview_select"
            )

            selected_file = next(
                (f for f in uploaded_files if f.name == preview_file), None
            )

            if selected_file:
                content = selected_file.getvalue()
                subtitles = SRTParser.parse(content)

                if subtitles:
                    # Show count
                    st.info(f"📊 Found **{len(subtitles)}** subtitles in this file")

                    # Show as dataframe
                    import pandas as pd
                    df = pd.DataFrame(subtitles)
                    df.columns = ['#', 'Start Time', 'End Time', 'Text']

                    # Limit preview to first 50
                    max_preview = 50
                    if len(subtitles) > max_preview:
                        st.warning(
                            f"Showing first {max_preview} of "
                            f"{len(subtitles)} subtitles"
                        )
                        df = df.head(max_preview)

                    st.dataframe(
                        df,
                        use_container_width=True,
                        height=400
                    )
                else:
                    st.warning("No subtitles found in this file.")

    else:
        # Empty state
        st.markdown("")
        st.markdown("")
        st.markdown("""
        <div style="text-align: center; padding: 3rem; color: #999;">
            <h2>👈 Upload SRT files to get started</h2>
            <p style="font-size: 1.1rem;">
                Drag and drop your .srt subtitle files in the upload area,
                <br>choose a format, and download your DOCX files instantly!
            </p>
            <br>
            <p style="font-size: 3rem;">📄 ➡️ 📝</p>
        </div>
        """, unsafe_allow_html=True)


# ── Footer ──
st.markdown("---")
st.markdown("""
<div style="text-align: center; color: #999; font-size: 0.8rem; padding: 1rem;">
    📝 SRT to DOCX Converter | Built with Streamlit |
    Upload → Convert → Download | No data stored on server
</div>
""", unsafe_allow_html=True)