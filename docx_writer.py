"""
DOCX Writer Module
Handles creation of Word documents from parsed subtitle data.
"""

from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


class DOCXWriter:
    """
    Creates formatted DOCX documents from subtitle data.

    Supports multiple output styles:
    - table: Organized table with columns
    - plain: Numbered entries with separators
    - formatted: Clean paragraphs with inline timestamps
    - text_only: Just subtitle text
    - script: Screenplay-style format
    """

    # Color constants
    COLOR_PRIMARY = RGBColor(0x1A, 0x47, 0x8A)
    COLOR_SECONDARY = RGBColor(0x5A, 0x9B, 0xD5)
    COLOR_TEXT = RGBColor(0x33, 0x33, 0x33)
    COLOR_MUTED = RGBColor(0x88, 0x88, 0x88)
    COLOR_LIGHT = RGBColor(0xCC, 0xCC, 0xCC)
    COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    COLOR_SUCCESS = RGBColor(0x27, 0xAE, 0x60)
    COLOR_TIMESTAMP = RGBColor(0x7F, 0x8C, 0x8D)

    # Available format styles
    AVAILABLE_STYLES = ['table', 'plain', 'formatted', 'text_only', 'script']

    def __init__(self):
        self.doc = None

    def create_document(self, subtitles, source_filename, output_path, style="table"):
        """
        Create a DOCX document from subtitle data.

        Args:
            subtitles: List of SubtitleEntry objects or dicts
            source_filename: Original SRT filename (for title)
            output_path: Path to save the DOCX file
            style: Output format style

        Raises:
            ValueError: If style is not recognized
        """
        if style not in self.AVAILABLE_STYLES:
            raise ValueError(
                f"Unknown style '{style}'. "
                f"Available: {', '.join(self.AVAILABLE_STYLES)}"
            )

        self.doc = Document()
        self._setup_page(self.doc)
        self._add_title(self.doc, source_filename)
        self._add_metadata(self.doc, subtitles, source_filename)
        self._add_spacer(self.doc)

        # Convert SubtitleEntry objects to dicts if needed
        subtitle_dicts = []
        for sub in subtitles:
            if hasattr(sub, 'to_dict'):
                subtitle_dicts.append(sub.to_dict())
            elif isinstance(sub, dict):
                subtitle_dicts.append(sub)
            else:
                raise TypeError(f"Unexpected subtitle type: {type(sub)}")

        # Write content based on selected style
        style_methods = {
            'table': self._write_table_format,
            'plain': self._write_plain_format,
            'formatted': self._write_formatted_format,
            'text_only': self._write_text_only_format,
            'script': self._write_script_format,
        }

        style_methods[style](self.doc, subtitle_dicts)

        # Add footer
        self._add_footer(self.doc, len(subtitle_dicts))

        # Save document
        self.doc.save(output_path)

    def _setup_page(self, doc):
        """Configure page margins and layout."""
        section = doc.sections[0]
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

    def _add_title(self, doc, source_filename):
        """Add document title."""
        title = doc.add_heading(level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Icon
        icon_run = title.add_run("📝 ")
        icon_run.font.size = Pt(20)

        # Title text
        title_run = title.add_run(f"Subtitles: {source_filename}")
        title_run.font.size = Pt(18)
        title_run.font.color.rgb = self.COLOR_PRIMARY
        title_run.bold = True

    def _add_metadata(self, doc, subtitles, source_filename):
        """Add document metadata section."""
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        info_parts = [
            f"Source: {source_filename}",
            f"Total Subtitles: {len(subtitles)}",
            f"Converted: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        ]

        # Calculate total duration if possible
        if subtitles:
            try:
                last_sub = subtitles[-1]
                if hasattr(last_sub, 'end_time'):
                    end_time = last_sub.end_time
                elif isinstance(last_sub, dict):
                    end_time = last_sub.get('end_time', '')
                else:
                    end_time = ''

                if end_time:
                    info_parts.append(f"Duration: ~{end_time}")
            except (AttributeError, IndexError):
                pass

        info_text = "  |  ".join(info_parts)
        meta_run = meta_para.add_run(info_text)
        meta_run.font.size = Pt(8)
        meta_run.font.color.rgb = self.COLOR_MUTED
        meta_run.italic = True

        # Add horizontal line
        line_para = doc.add_paragraph()
        line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line_run = line_para.add_run("━" * 80)
        line_run.font.size = Pt(6)
        line_run.font.color.rgb = self.COLOR_LIGHT

    def _add_spacer(self, doc):
        """Add a blank spacer paragraph."""
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(2)
        spacer.paragraph_format.space_after = Pt(2)

    def _add_footer(self, doc, count):
        """Add document footer."""
        self._add_spacer(doc)

        line_para = doc.add_paragraph()
        line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line_run = line_para.add_run("━" * 80)
        line_run.font.size = Pt(6)
        line_run.font.color.rgb = self.COLOR_LIGHT

        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run(
            f"Generated by SRT to DOCX Converter  |  "
            f"{count} subtitle entries  |  "
            f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        )
        footer_run.font.size = Pt(7)
        footer_run.font.color.rgb = self.COLOR_MUTED
        footer_run.italic = True

    # ─────────────────────────────────────────────
    # FORMAT: TABLE
    # ─────────────────────────────────────────────
    def _write_table_format(self, doc, subtitles):
        """Write subtitles in a structured table format."""
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        # Header row
        headers = ['#', 'Start Time', 'End Time', 'Subtitle Text']
        header_cells = table.rows[0].cells

        for i, header_text in enumerate(headers):
            cell = header_cells[i]
            cell.text = ''
            para = cell.paragraphs[0]
            run = para.add_run(header_text)
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = self.COLOR_WHITE
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Data rows
        for sub in subtitles:
            row_cells = table.add_row().cells

            # Index
            idx_para = row_cells[0].paragraphs[0]
            idx_run = idx_para.add_run(str(sub['index']))
            idx_run.font.size = Pt(9)
            idx_run.font.color.rgb = self.COLOR_PRIMARY
            idx_run.bold = True
            idx_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Start Time
            start_para = row_cells[1].paragraphs[0]
            start_run = start_para.add_run(sub['start_time'])
            start_run.font.size = Pt(8)
            start_run.font.color.rgb = self.COLOR_TIMESTAMP
            start_run.font.name = 'Consolas'
            start_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # End Time
            end_para = row_cells[2].paragraphs[0]
            end_run = end_para.add_run(sub['end_time'])
            end_run.font.size = Pt(8)
            end_run.font.color.rgb = self.COLOR_TIMESTAMP
            end_run.font.name = 'Consolas'
            end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Text
            text_para = row_cells[3].paragraphs[0]
            text_run = text_para.add_run(sub['text'])
            text_run.font.size = Pt(10)
            text_run.font.color.rgb = self.COLOR_TEXT

        # Set column widths
        for row in table.rows:
            row.cells[0].width = Cm(1.2)
            row.cells[1].width = Cm(3.5)
            row.cells[2].width = Cm(3.5)
            row.cells[3].width = Cm(10.0)

    # ─────────────────────────────────────────────
    # FORMAT: PLAIN
    # ─────────────────────────────────────────────
    def _write_plain_format(self, doc, subtitles):
        """Write subtitles in plain numbered format with separators."""
        for i, sub in enumerate(subtitles):
            # Header line: index + timestamp
            header_para = doc.add_paragraph()
            header_para.paragraph_format.space_before = Pt(6)
            header_para.paragraph_format.space_after = Pt(2)

            # Index badge
            index_run = header_para.add_run(f" {sub['index']} ")
            index_run.bold = True
            index_run.font.size = Pt(9)
            index_run.font.color.rgb = self.COLOR_WHITE

            # Separator
            sep_run = header_para.add_run("  ")
            sep_run.font.size = Pt(9)

            # Timestamp
            time_run = header_para.add_run(
                f"⏱ {sub['start_time']}  →  {sub['end_time']}"
            )
            time_run.font.size = Pt(9)
            time_run.font.color.rgb = self.COLOR_TIMESTAMP
            time_run.italic = True
            time_run.font.name = 'Consolas'

            # Subtitle text
            text_para = doc.add_paragraph()
            text_para.paragraph_format.space_before = Pt(0)
            text_para.paragraph_format.space_after = Pt(4)
            text_para.paragraph_format.left_indent = Inches(0.3)

            text_run = text_para.add_run(sub['text'])
            text_run.font.size = Pt(11)
            text_run.font.color.rgb = self.COLOR_TEXT
            text_run.font.name = 'Calibri'

            # Separator line (except for last entry)
            if i < len(subtitles) - 1:
                sep_para = doc.add_paragraph()
                sep_para.paragraph_format.space_before = Pt(2)
                sep_para.paragraph_format.space_after = Pt(2)
                sep_line = sep_para.add_run("─" * 70)
                sep_line.font.size = Pt(5)
                sep_line.font.color.rgb = self.COLOR_LIGHT

    # ─────────────────────────────────────────────
    # FORMAT: FORMATTED
    # ─────────────────────────────────────────────
    def _write_formatted_format(self, doc, subtitles):
        """Write subtitles as formatted paragraphs with inline timestamps."""
        for sub in subtitles:
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.line_spacing = Pt(16)

            # Timestamp in brackets
            time_run = para.add_run(
                f"[{sub['start_time']} – {sub['end_time']}]  "
            )
            time_run.font.size = Pt(8)
            time_run.font.color.rgb = self.COLOR_TIMESTAMP
            time_run.italic = True
            time_run.font.name = 'Consolas'

            # Subtitle text
            text_run = para.add_run(sub['text'])
            text_run.font.size = Pt(11)
            text_run.font.color.rgb = self.COLOR_TEXT
            text_run.font.name = 'Calibri'

    # ─────────────────────────────────────────────
    # FORMAT: TEXT ONLY
    # ─────────────────────────────────────────────
    def _write_text_only_format(self, doc, subtitles):
        """Write only subtitle text without any timestamps or numbers."""
        for sub in subtitles:
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after = Pt(6)
            para.paragraph_format.line_spacing = Pt(18)

            text_run = para.add_run(sub['text'])
            text_run.font.size = Pt(12)
            text_run.font.color.rgb = self.COLOR_TEXT
            text_run.font.name = 'Georgia'

    # ─────────────────────────────────────────────
    # FORMAT: SCRIPT
    # ─────────────────────────────────────────────
    def _write_script_format(self, doc, subtitles):
        """Write subtitles in a screenplay/script style."""
        for sub in subtitles:
            # Timestamp as scene marker
            time_para = doc.add_paragraph()
            time_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            time_para.paragraph_format.space_before = Pt(12)
            time_para.paragraph_format.space_after = Pt(2)

            marker_run = time_para.add_run(
                f"[{sub['start_time']} → {sub['end_time']}]"
            )
            marker_run.font.size = Pt(8)
            marker_run.font.color.rgb = self.COLOR_SECONDARY
            marker_run.bold = True
            marker_run.font.name = 'Courier New'

            # Dialogue text (centered, screenplay style)
            text_para = doc.add_paragraph()
            text_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            text_para.paragraph_format.space_before = Pt(0)
            text_para.paragraph_format.space_after = Pt(4)
            text_para.paragraph_format.left_indent = Inches(1.5)
            text_para.paragraph_format.right_indent = Inches(1.5)

            text_run = text_para.add_run(sub['text'])
            text_run.font.size = Pt(11)
            text_run.font.color.rgb = self.COLOR_TEXT
            text_run.font.name = 'Courier New'